import PIL
from PIL import Image
from django.core.files.uploadedfile import InMemoryUploadedFile
from rest_framework.generics import CreateAPIView
from rest_framework.response import Response
from rest_framework import status
from .models import UserInput, UserInputGenai, ChatGptModel
from .serializer import UserInputCohereSerializer, UserInputGeminiSerializer, GptSerializer
import cohere
import google.generativeai as genai
from docx import Document
from PyPDF2 import PdfReader
from os.path import splitext

co = cohere.Client('APIKEYY')

prompt = """Act as a Leadership Management Tutor working at Desklib (online study resources and Homework Help website). 
        Give an answer to the question asked in very brief, strictly under 300 words. Make sure to only answer if you are confident; otherwise, 
        recommend me on using Desklib Expert Help services to get a step-by-step solution with explanation."""


class UserInputCohereAPIView(CreateAPIView):
    queryset = UserInput.objects.all()
    serializer_class = UserInputCohereSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        docs = validated_data['document']
        if docs is None:
            user_message = validated_data.get('user_message', '')
        else:
            file_name, file_extension = splitext(docs.name.lower())
            if file_extension == '.docx':
                doc = Document(docs)
                user_message = validated_data.get('user_message', '')
                for paragraph in doc.paragraphs:
                    content = paragraph.text
                    user_message += '\n' + content
            if file_extension == '.pdf':
                reader = PdfReader(docs)
                page = reader.pages[0]
                text = page.extract_text()
                user_message = validated_data.get('user_message', '')
                user_message += '\n' + text

        conversation = [
            {"role": "system", "content": prompt, "message": prompt},
            {"role": "user", "content": user_message, "message": user_message}
        ]
        response = co.chat(
            chat_history=conversation,
            message=user_message,
            connectors=[{"id": "web-search"}]
        )
        validated_data['message'] = response.text
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)


genai.configure(api_key="APIKEYY")


class UserInputGeminiCreateAPIView(CreateAPIView):
    queryset = UserInputGenai.objects.all()
    serializer_class = UserInputGeminiSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        user_message = validated_data.get('user_message', '')
        # Set up the model
        generation_config = {
            "temperature": 0.9,
            "top_p": 1,
            "top_k": 1,
            "max_output_tokens": 400,
        }

        safety_settings = [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            }
        ]
        if validated_data.get('image') is None:
            docs = validated_data['document']
            if docs:
                file_name, file_extension = splitext(docs.name.lower())
                if file_extension == '.docx':
                    doc = Document(docs)
                    for paragraph in doc.paragraphs:
                        content = paragraph.text
                if file_extension == '.pdf':
                    reader = PdfReader(docs)
                    page = reader.pages[0]
                    text = page.extract_text()
                    user_message = validated_data.get('user_message', '')
                    user_message += '\n' + text
                    content = user_message
            model = genai.GenerativeModel(model_name="gemini-pro",
                                          generation_config=generation_config,
                                          safety_settings=safety_settings)
            if docs:
                prompt_parts = [prompt + user_message +'\n' + content]
            else:
                prompt_parts = [prompt + '\n' + user_message]
            response = model.generate_content(prompt_parts)
            validated_data['bot_message'] = response.text
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)
        if validated_data.get('image') is not None and validated_data.get('user_message') is not None:
            image = validated_data.get('image')
            model = genai.GenerativeModel(model_name="gemini-pro-vision",
                                          generation_config=generation_config,
                                          safety_settings=safety_settings)
            img = PIL.Image.open(image)
            response = model.generate_content([prompt + '\n' + user_message, img], stream=True)
            response.resolve()
            validated_data['image'] = image
            validated_data['bot_message'] = response.text
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)

        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)


from openai import OpenAI
import openai
import base64

client = openai.Client(api_key="APIKEYY")

def encode_image(image):
    if isinstance(image, InMemoryUploadedFile):
        image_content = image.read()
        return base64.b64encode(image_content).decode('utf-8')
    else:
        raise TypeError("Invalid image format")

class ChatGpt4Api(CreateAPIView):
    queryset = ChatGptModel.objects.all()
    serializer_class = GptSerializer

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        user_message = validated_data.get('user_message', '')
        image = validated_data.get('image')
        if image:
            base64_image = encode_image(image)

            response = client.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=[
                    {
                        "role": "system",
                        "content": prompt,
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text":user_message},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                    "detail": "high"
                                },
                            },
                        ],
                    }
                ],
                max_tokens=300,
            )
        else:
            conversation = [
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_message},
            ]
            response = client.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=conversation,
                max_tokens=300,
            )
        gpt3_response = response.choices[0].message.content
        validated_data['bot_message'] = gpt3_response
        chat_gpt_instance = ChatGptModel.objects.create(
            user_message=user_message,
            bot_message=gpt3_response,
            image=validated_data['image'],
        )
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
